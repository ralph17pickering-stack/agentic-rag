from backend.app.services.extraction import extract_text


# ── HTML extraction ────────────────────────────────────────────────────────────

def test_html_heading_h1_becomes_markdown():
    html = b"<h1>Hello World</h1>"
    result = extract_text(html, "html")
    assert "# Hello World" in result


def test_html_heading_h2_becomes_markdown():
    html = b"<h2>Sub Section</h2>"
    result = extract_text(html, "html")
    assert "## Sub Section" in result


def test_html_heading_h3_becomes_markdown():
    html = b"<h3>Deep</h3>"
    result = extract_text(html, "html")
    assert "### Deep" in result


def test_html_heading_h4_becomes_markdown():
    html = b"<h4>Deeper</h4>"
    result = extract_text(html, "html")
    assert "#### Deeper" in result


def test_html_list_items_become_dashes():
    html = b"<ul><li>Alpha</li><li>Beta</li></ul>"
    result = extract_text(html, "html")
    assert "- Alpha" in result
    assert "- Beta" in result


def test_html_table_becomes_markdown_table():
    html = b"""
    <table>
      <tr><th>Name</th><th>Age</th></tr>
      <tr><td>Alice</td><td>30</td></tr>
    </table>
    """
    result = extract_text(html, "html")
    assert "| Name | Age |" in result
    assert "| --- | --- |" in result
    assert "| Alice | 30 |" in result


def test_html_script_tags_stripped():
    html = b"<p>Keep this</p><script>var x = 1;</script>"
    result = extract_text(html, "html")
    assert "Keep this" in result
    assert "var x" not in result


def test_html_nav_stripped():
    html = b"<nav><a>Menu</a></nav><p>Content</p>"
    result = extract_text(html, "html")
    assert "Content" in result
    assert "Menu" not in result


def test_html_preserves_paragraph_text():
    html = b"<p>Some body text here.</p>"
    result = extract_text(html, "html")
    assert "Some body text here." in result


def test_html_no_excessive_blank_lines():
    html = b"<p>A</p><p>B</p><p>C</p>"
    result = extract_text(html, "html")
    # Should not have 3+ consecutive newlines
    assert "\n\n\n" not in result


def test_html_ordered_list_items_numbered():
    html = b"<ol><li>First</li><li>Second</li></ol>"
    result = extract_text(html, "html")
    assert "1. First" in result
    assert "2. Second" in result


# ── DOCX extraction ────────────────────────────────────────────────────────────

from io import BytesIO
from docx import Document as DocxDocument


def _make_docx_bytes(**kwargs) -> bytes:
    """Create a minimal .docx in memory.
    kwargs:
      paragraphs: list of (text, style_name) tuples
      tables: list of rows, each row is a list of cell strings
    """
    doc = DocxDocument()
    for text, style in kwargs.get("paragraphs", []):
        doc.add_paragraph(text, style=style)
    for table_data in kwargs.get("tables", []):
        t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_text in enumerate(row):
                t.cell(r_idx, c_idx).text = cell_text
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_heading1_becomes_markdown():
    docx_bytes = _make_docx_bytes(paragraphs=[("Introduction", "Heading 1")])
    result = extract_text(docx_bytes, "docx")
    assert "# Introduction" in result


def test_docx_heading2_becomes_markdown():
    docx_bytes = _make_docx_bytes(paragraphs=[("Sub-topic", "Heading 2")])
    result = extract_text(docx_bytes, "docx")
    assert "## Sub-topic" in result


def test_docx_body_paragraph_has_no_prefix():
    docx_bytes = _make_docx_bytes(paragraphs=[("Just a sentence.", "Normal")])
    result = extract_text(docx_bytes, "docx")
    assert "Just a sentence." in result
    assert "# Just a sentence." not in result


def test_docx_table_becomes_markdown_table():
    docx_bytes = _make_docx_bytes(
        tables=[[["Name", "Score"], ["Alice", "95"], ["Bob", "87"]]]
    )
    result = extract_text(docx_bytes, "docx")
    assert "| Name | Score |" in result
    assert "| --- | --- |" in result
    assert "| Alice | 95 |" in result
