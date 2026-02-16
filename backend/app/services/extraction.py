"""Text extraction dispatcher — converts any supported file format to plain text."""

import csv
from io import BytesIO, StringIO

import pypdf
from bs4 import BeautifulSoup
from docx import Document as DocxDocument


SUPPORTED_TYPES = {"txt", "md", "pdf", "docx", "csv", "html"}


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extract plain text from file bytes based on file_type."""
    if file_type in ("txt", "md"):
        return file_bytes.decode("utf-8")

    if file_type == "pdf":
        return _extract_pdf(file_bytes)

    if file_type == "docx":
        return _extract_docx(file_bytes)

    if file_type == "csv":
        return _extract_csv(file_bytes)

    if file_type == "html":
        return _extract_html(file_bytes)

    raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    reader = pypdf.PdfReader(BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_csv(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8")
    reader = csv.DictReader(StringIO(text))
    rows = []
    for i, row in enumerate(reader, start=1):
        pairs = ", ".join(f"{k}={v}" for k, v in row.items() if v)
        rows.append(f"Row {i}: {pairs}")
    return "\n".join(rows)


_STRIP_TAGS = {"script", "style", "nav", "header", "footer", "aside"}
_HEADING_MAP = {"h1": "#", "h2": "##", "h3": "###", "h4": "####"}


def _extract_html(file_bytes: bytes) -> str:
    import re
    soup = BeautifulSoup(file_bytes, "html.parser")

    # Remove boilerplate elements
    for tag in soup.find_all(_STRIP_TAGS):
        tag.decompose()

    lines: list[str] = []

    def _walk(node) -> None:
        from bs4 import NavigableString, Tag
        if isinstance(node, NavigableString):
            text = node.strip()
            if text:
                lines.append(text)
            return

        tag_name = node.name
        if tag_name is None:
            return

        if tag_name in _HEADING_MAP:
            prefix = _HEADING_MAP[tag_name]
            text = node.get_text(" ", strip=True)
            if text:
                lines.append(f"\n{prefix} {text}")
            return

        if tag_name == "li":
            text = node.get_text(" ", strip=True)
            if text:
                lines.append(f"- {text}")
            return

        if tag_name == "table":
            lines.append(_table_to_markdown(node))
            return

        if tag_name in ("ul", "ol"):
            for child in node.children:
                _walk(child)
            return

        if tag_name == "p":
            text = node.get_text(" ", strip=True)
            if text:
                lines.append(text)
            return

        # Generic: recurse into children
        for child in node.children:
            _walk(child)

    _walk(soup.body or soup)

    # Collapse excessive blank lines
    result = "\n".join(lines)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def _table_to_markdown(table_node) -> str:
    """Convert a BeautifulSoup table element to GFM markdown table."""
    rows = table_node.find_all("tr")
    if not rows:
        return ""

    table_lines: list[str] = []
    header_done = False

    for row in rows:
        cells = [c.get_text(" ", strip=True) for c in row.find_all(["th", "td"])]
        if not cells:
            continue
        table_lines.append("| " + " | ".join(cells) + " |")
        if not header_done:
            table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            header_done = True

    return "\n".join(table_lines)
